"""Build the PRDCP v2 developer guide from the active implementation contract."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "PRDCP_v2_Developer_Guide.docx"

NAVY = "18324A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "68737D"
DARK = "17202A"
WHITE = "FFFFFF"
GREEN = "2F6B4F"
GOLD = "7A5A00"
RED = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 font: str = "Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Yu Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_page_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, MID_GRAY, 0, 16),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(DARK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.10)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(section, *, first: bool = False) -> None:
    section.different_first_page_header_footer = first
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("PRDCP v2  |  Developer Guide")
    set_run_font(run, size=8.5, color=MID_GRAY, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("PRDCP v2  •  ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    add_page_field(p, "PAGE")
    run = p.add_run(" / ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    add_page_field(p, "NUMPAGES")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_para(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item))


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_run_font(p.add_run(item))


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    set_run_font(p.add_run(text), size=8.5, font="Consolas")


def add_note(doc: Document, label: str, text: str, *, color: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    lead = p.add_run(label + "  ")
    set_run_font(lead, bold=True, color=NAVY)
    set_run_font(p.add_run(text))
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[int], *, font_size: float = 9.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if len(value) < 18 and index == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(value), size=font_size)
    set_table_geometry(table, widths)
    for row in table.rows:
        prevent_row_split(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_chapter(doc: Document, number: int, title: str) -> None:
    add_heading(doc, f"{number}. {title}", 1)


def build() -> Path:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    # Editorial cover for a long-form technical manual.
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run("DEVELOPER REFERENCE"), size=10, color=GOLD, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("PRDCP v2"), size=30, color=NAVY, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("Developer Guide"), size=17, color=DARK_BLUE, bold=True)
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_before = Pt(12)
    desc.paragraph_format.space_after = Pt(30)
    set_run_font(
        desc.add_run("5 Layer Architecture • PMP v2 • Revision v1 • Recovery • Delivery"),
        size=11,
        color=MID_GRAY,
    )
    add_note(
        doc,
        "対象",
        "PRDCPを初めて保守・拡張する開発者。active codeを正本とし、旧checkpointの非破壊互換を含む。",
        color=LIGHTER_BLUE,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    set_run_font(meta.add_run(f"Version 2.0.0  |  更新日 {date.today().isoformat()}"), size=10, color=MID_GRAY)
    doc.add_page_break()

    add_heading(doc, "Guide Map", 1)
    add_para(doc, "本書は30章を、全体像、実行基盤、安全境界、Revision、運用・拡張の順に構成しています。")
    add_table(
        doc,
        ["Part", "Chapters", "読む目的"],
        [
            ["A", "1–8", "設計思想、5 Layer、Agent、Runtime、PMP、RD、Promptを理解する"],
            ["B", "9–17", "Provider、Structured Output、Retrieval、Storage、Outbox、State、Human/Final Gateを理解する"],
            ["C", "18–21", "共通Revision、Layer別往復、Recovery/Retry/Repair、Safe Modeを理解する"],
            ["D", "22–30", "CLI/Discord、Delivery、監査、テスト、拡張手順、制約、障害対応を実践する"],
        ],
        [900, 1350, 7110],
    )
    add_note(doc, "Canonical boundary", "Layer内はin-process direct call、Layer間はfile Outboxです。PMPはprotocol/envelopeでありmessage brokerではありません。")
    doc.add_page_break()

    add_chapter(doc, 1, "PRDCP v2概要")
    add_para(doc, "PRDCP v2は、社会で共有される一般論をTopic探索、Evidence収集、多角的審議、結論選択、動画台本制作へ分解する5層マルチエージェントシステムです。LangGraph/LangChainに依存せず、Manager、PMP、file Outbox、JSON checkpointでオーケストレーションを明示しています。")
    add_bullets(doc, [
        "31 Role Definition、31 Message Type、7 PMP Status、6 Layer間Handoff契約を起動時に検証する。",
        "OpenRouterとMockを同じProvider abstractionで扱い、22 root Structured Output SchemaをFail Closedで監査する。",
        "Human Evidence GateとHuman Selection Gateを自動推測せず、Deliveryまでの判断履歴を保存する。",
        "既存の正常なArtifact、Retrieval Context、Reservation、CheckpointをRecovery時に再利用する。",
    ])
    add_code(doc, "Entry point: main.py\nComposition root: runtime.py\nCanonical storage: storage/data/\nVerification: scripts/verify.py")

    add_chapter(doc, 2, "設計思想")
    add_para(doc, "設計の中心は、品質と費用と復旧可能性を同時に守ることです。Managerは成果物を生成するだけでなく、実行前の境界検証、保存、依存閉包、予約、Human Gate、再開位置を管理します。")
    add_table(doc, ["原則", "実装上の意味"], [
        ["Fail Closed", "Schema、PMP、hash、route、budget、Human Gateが不明なら続行しない"],
        ["Checkpoint First", "Revision planとtask identityをProvider呼び出し前に保存する"],
        ["Exactly Once Intent", "create-once Artifactとpersistent reservationで曖昧な二重送信を防ぐ"],
        ["Typed Provenance", "claim→analysis→evidence→sourceを型別IDで追跡する"],
        ["Minimal Replay", "完了済み高コストstageを再実行せず、必要な依存閉包だけを更新する"],
        ["Human Authority", "Evidence受容、再調査実行、最終候補選択を別々の明示判断として保存する"],
    ], [1800, 7560])

    add_chapter(doc, 3, "5 Layer Architecture")
    add_table(doc, ["Producer", "Researcher", "Deliberation", "Conclusion", "Playwright"], [[
        "Topic・Plan", "Evidence・Report", "Analysis・Integration", "Candidates・Selection", "Script・Delivery"
    ]], [1872, 1872, 1872, 1872, 1872], font_size=8.7)
    add_para(doc, "通常のforward workflowは左から右へ進みます。品質不足は同一LayerのInternal Revision、または直前LayerへのUpstream Revisionとして右から左へ一段だけ戻ります。Layer skippingは共通Validatorが拒否します。")
    add_table(doc, ["Layer", "入力", "出力", "停止境界"], [
        ["Producer", "User topic", "research_plan", "Quality Review"],
        ["Researcher", "research_plan", "research_result", "Human Evidence Gate"],
        ["Deliberation", "research_result", "deliberation_result", "Quality/Readiness"],
        ["Conclusion", "deliberation_result", "conclusion_handoff", "Human Selection"],
        ["Playwright", "conclusion_handoff", "final_script_delivery", "Final Gate"],
    ], [1350, 2250, 2520, 3240])

    add_chapter(doc, 4, "Agent構成と責務")
    add_table(doc, ["Layer", "RD数", "Agent / Manager"], [
        ["Producer", "6", "manager, topic_scout, topic_selector, general_opinion_analyst, research_planner, quality_reviewer"],
        ["Researcher", "9", "manager, academic, government, industry, expert, politician, news, public_opinion, quality_reviewer"],
        ["Deliberation", "6", "manager, argument, causal_structural, stakeholder_response, counterargument, quality_reviewer"],
        ["Conclusion", "5", "manager, position_generator, decision_evaluator, decision_integrator, quality_reviewer"],
        ["Playwright", "5", "manager, narrative_architect, scriptwriter, evidence_citation_editor, visual_director"],
    ], [1380, 660, 7320], font_size=8.8)
    add_note(doc, "Manager ownership", "Manager-owned integrationやFinal Gateには独立Agent PMPがない箇所があります。その場合もmanager RD usage、checkpoint trace、Revision auditで実行根拠を追跡します。")

    add_chapter(doc, 5, "Runtime / Orchestrator")
    add_para(doc, "`runtime.py`はSettings、Provider、RetrievalCoordinator、RoleDefinitionLoader、Layer Repository、Registry、Managerを組み立てるcomposition rootです。5 Managerは同じ実効Provider/Safe Mode設定とRD snapshot生成規則を共有します。")
    add_table(doc, ["主要関数", "責務"], [
        ["runtime.build_provider()", "mock/openrouter選択とprovider reservation root設定"],
        ["runtime.build_retrieval_coordinator()", "検索Provider、retrieval reservation、Context保存境界を構築"],
        ["runtime.build_role_definition_loader()", "31 RDをSTRICT preloadしaccess logを設定"],
        ["runtime.build_all_managers()", "1 Provider、1 RD Loaderを5 Layerへ共有して構築"],
        ["cli_app.commands.build_*", "CLI引数の実効設定をManagerへ適用"],
    ], [3300, 6060])

    add_chapter(doc, 6, "PMP v2")
    add_para(doc, "PMPMessageは通信内容を包む監査可能なenvelopeです。`common/models/pmp.py`がPydantic正本、`specifications/common/`が機械可読registryです。")
    add_table(doc, ["要素", "意味"], [
        ["workflow_id", "5 Layerを通じて保存するUUID"],
        ["message_id / parent_message_id", "request、response、revisionの親子相関"],
        ["sender_agent_id / receiver_agent_id", "Registryで検証する送受信主体"],
        ["message_type", "task/result/revision_request/revision_result/handoff等の31種"],
        ["status", "pending/running/completed/revision_required/error等の有限集合"],
        ["context / routing / metadata", "stage、execution order、revision target、retry、RD hashを保存"],
    ], [2700, 6660])
    add_note(doc, "Transport", "PMPはbrokerではありません。Layer内は直接call、Layer間だけfile Outboxです。DiscordはControl Planeで、PMPを配送しません。")

    add_chapter(doc, 7, "RD / RD Loader")
    add_para(doc, "31 JSON Role Definitionは`role_definitions/registry.json`から解決されます。LoaderはSchema/agent ID/version/timeout/message type/boundaryを検証し、1件でも不正ならSTRICT起動を停止します。")
    add_code(doc, "common/role_definitions/loader.py       RoleDefinitionLoader\ncommon/role_definitions/validator.py    schema/boundary validation\ncommon/role_definitions/agent_runtime.py runtime extraction\nrole_definitions/<layer>/*.json          canonical RD\nstorage/data/logs/rd_access.jsonl        access trace")
    add_bullets(doc, [
        "Agent実行ごとに固定snapshotを使い、途中のRD変更を同一callへ混入させない。",
        "LLMへ渡すRole ContextとRuntime timeout/revision limit/model設定を分離する。",
        "応答PMPにRD ID、version、SHA-256を保存し、Manager利用もstateへ記録する。",
        "reload時の不正更新は旧cacheへ黙ってfallbackせずFail Closedにする。",
    ])

    add_chapter(doc, 8, "Prompt Builder")
    add_para(doc, "`common/prompting/prompt_builder.py`はsystem rules、RD role context、task input、Structured Output指示を決定的な順序で合成します。Schema全文をsystem promptへ二重埋め込みせず、Providerの`response_format`を権威ある境界とします。")
    add_table(doc, ["優先順", "内容"], [
        ["1", "システム安全規則とPMP/ID/証拠境界"],
        ["2", "検証済みRD snapshotの役割・禁止事項"],
        ["3", "workflow/task固有の入力と保存済みContext view"],
        ["4", "単一JSON objectとして応答するStrict Output指示"],
    ], [1100, 8260])

    add_chapter(doc, 9, "Provider abstraction")
    add_para(doc, "`providers/`はMockとOpenRouterを同じAgent実行pipelineへ接続します。Provider callはlogical task ID、workflow ID、Agent、model、output schema、reservationへ結び付きます。")
    add_bullets(doc, [
        "OpenRouterは`provider.require_parameters=true`とstrict `response_format`を弱めない。",
        "公開Endpoint metadataでmodel/alias/structured output能力を予約前に検査する。",
        "曖昧な通信障害は自動retryせず、operator retryを別task identityで一回だけ許可する。",
        "異なるmodelのcontract/capability repairは専用authorizationと互換bindingを持つ。",
    ])
    add_code(doc, "providers/openrouter_provider.py\nproviders/mock_provider.py\ncommon/provider_model_compatibility.py\nstorage/data/provider_call_reservations/")

    add_chapter(doc, 10, "Structured Output")
    add_para(doc, "`common/structured_outputs.py`はPydantic Schemaをコピーし、strict normalization、request固有specialization、再帰validationの順で処理します。22 root schemaすべてを同じ監査対象にします。")
    add_table(doc, ["検査対象", "必須条件"], [
        ["root / nested object", "additionalProperties=false"],
        ["propertiesを持つobject", "required集合とproperty key集合が完全一致"],
        ["$defs", "参照先objectを同じ規則で再帰検査"],
        ["array.items", "items以下のobject/unionを再帰検査"],
        ["anyOf / union", "全branchを検査し未解決$refや$ref siblingを拒否"],
        ["free-form dict", "意味を壊して閉じず、Structured Output境界では明示modelを要求"],
    ], [2700, 6660])
    add_note(doc, "Dynamic constraints", "Retrieval Agentは長いtitle/URL/excerptではなく短いsource_idだけをenum拘束し、保存Contextからidentityを決定論的に復元します。")

    add_chapter(doc, 11, "Retrieval")
    add_para(doc, "検索とStructured Reasoningは別工程です。General Opinion Analystと7 Researcher specialistだけがRetrievalを使い、Research Plannerには付与しません。")
    add_steps(doc, [
        "検索task identityとretrieval reservationを作成する。",
        "URL、title、excerpt、取得時刻をretrieval_contextsへ永続化する。",
        "Context hashを固定し、短いruntime viewをReasoning Agentへ渡す。",
        "LLM出力のsource_idを保存集合へ照合し、identity/excerptを決定論的にhydrateする。",
        "Reasoning失敗時は保存Contextを再利用し、検索を自動再実行しない。",
    ])
    add_code(doc, "retrieval/coordinator.py\nstorage/data/retrieval_contexts/<workflow_id>/\nstorage/data/retrieval_call_reservations/")

    add_chapter(doc, 12, "Storage")
    add_para(doc, "`storage/data`はRuntimeのcanonical storageです。既存JSONを破壊的migrationせず、新フィールドはPydantic defaultとread adapterで補います。書込みは可能な限りatomic、create-once、append-onlyに分けます。")
    add_table(doc, ["領域", "例"], [
        ["Workflow state", "workflows/<layer>/<workflow_id>.json と messages.jsonl"],
        ["Domain artifacts", "research_reports, deliberation_results, conclusion_packages, script_drafts"],
        ["Outbox", "forward handoff、revision_requests、revision_results"],
        ["Safety ledgers", "provider/retrieval reservations、authorizations、revision_budget"],
        ["Audit", "revision_audit、integrity repair、deterministic repair、RD access"],
        ["Delivery", "deliveries/<workflow_id>/ の6成果物"],
    ], [2500, 6860])

    add_chapter(doc, 13, "Outbox")
    add_para(doc, "Layer間HandoffとUpstream Revisionはfile Outboxで配送します。通常forward handoffはworkflow単位、共通Revisionはrequest単位です。")
    add_code(doc, "outbox/researcher/<workflow_id>.json\noutbox/deliberation/<workflow_id>.json\noutbox/conclusion/<workflow_id>.json\noutbox/playwright/<workflow_id>.json\noutbox/revision_requests/<target_layer>/<workflow_id>/<request_id>.json\noutbox/revision_results/<requester_layer>/<workflow_id>/<request_id>.json")
    add_bullets(doc, [
        "Writerはcreate-once。既存identityと異なるpayloadの上書きを拒否する。",
        "Consumerはworkflow、request、parent message、finding、Artifact hashを照合する。",
        "Internal RevisionはOutboxへ出さず、内部監査Artifactとして保存する。",
        "旧単一ファイルOutboxはread adapterで読み、新形式への強制書換えは行わない。",
    ])

    add_chapter(doc, 14, "Workflow State")
    add_para(doc, "各Layer stateは成果物、message history、checkpoint、error、revision履歴に加え、additiveな`RevisionControlState`を持ちます。旧checkpointにフィールドがなくても`idle`で読めます。")
    add_table(doc, ["Revision phase", "意味"], [
        ["idle", "Revisionなし、または旧stateの互換default"],
        ["planned / requested", "対象、finding、base artifactを固定済み"],
        ["authorization_required", "有料実行の明示承認待ち"],
        ["waiting_upstream_result", "隣接上流Layerの相関Result待ち"],
        ["consuming_request / executing", "Request検証または最小依存閉包を実行中"],
        ["result_ready / completed", "Result保存または下流反映が完了"],
        ["blocked", "budget、stale、correlation、Human Gate等で安全停止"],
    ], [2700, 6660])

    add_chapter(doc, 15, "Human Evidence Gate")
    add_para(doc, "Researcher Quality Review後は判定にかかわらず人間境界で停止します。Evidence GapとHard Integrity Failureを混同せず、Human Decisionと実行authorizationも分離します。")
    add_table(doc, ["Decision", "許可条件と結果"], [
        ["ACCEPT", "未解決Evidence Gapがない場合に下流Handoffを許可"],
        ["ACCEPT_WITH_LIMITATIONS", "GapをEvidence化せずlimitationsとして全下流へ伝播"],
        ["REVISE", "0-callでRevision Planを保存。別のrevision-execute承認までProviderを呼ばない"],
    ], [2500, 6860])
    add_note(doc, "Hard Integrity", "Schema/PMP/provenance/矛盾したsource relationはHuman Decisionでoverrideできません。allowlist型deterministic repairでのみ修復し、解けなければFail Closedです。", color="FCE8E6")

    add_chapter(doc, 16, "Human Selection Gate")
    add_para(doc, "Conclusion Quality Gateを通過してもFinal Conclusionは自動確定しません。`WAITING_HUMAN_SELECTION`で候補・評価・limitationsを提示し、明示選択または明示統合だけを保存します。")
    add_bullets(doc, [
        "同じselectionはidempotent、異なる再選択は契約上許可されたRevision後だけ。",
        "Playwrightの構造修正は`unchanged`を宣言して既存selectionを保持できる。",
        "意味変更は`reselection_required`を宣言し、旧selectionとFinal Conclusionを無効化して再選択で停止する。",
        "Human Selection PMPとFinal Conclusion hashをPlaywright開始前に検証する。",
    ])

    add_chapter(doc, 17, "Final Gate")
    add_para(doc, "Playwrightは独立Quality Reviewerを置かず、deterministic Validator、Evidence & Citation Editor、Manager Final Gateで納品可能性を判定します。")
    add_table(doc, ["Gate", "代表検査"], [
        ["Identity", "Final Conclusion ID/hash、Human Selection、workflow一致"],
        ["Claim", "Script claim集合とManifest supported_claim_idsの完全一致"],
        ["Citation", "引用必須Paragraph→Claim→Evidence→Source mappingとlocator"],
        ["Disclosure", "accepted gapとlimitationsをEvidenceへ昇格させず開示"],
        ["Visual", "section/paragraph/evidence/source/asset参照とchart出典"],
        ["Delivery", "blocking finding 0、6ファイル、Delivery message exactly once"],
    ], [2300, 7060])

    add_chapter(doc, 18, "Revision Architecture")
    add_para(doc, "`common/models/revision.py`、`common/validation/revision_validator.py`、`storage/revision_exchange_repository.py`が共通基盤です。Request/Result、budget、authorization、audit、stateをLayer固有Schemaから分離します。")
    add_table(doc, ["Contract", "重要フィールド"], [
        ["RevisionRequestV1", "request/root/parent ID、workflow、route、source/target layer、epoch、review/finding、target Agent、base Artifact hash、actions、acceptance、retrieval許可、idempotency"],
        ["RevisionResultV1", "request/message相関、producer/requester layer、status、base/result Artifact hash、finding disposition、selection impact、reservation/call数、idempotency"],
        ["RevisionControlState", "phase、epoch、active request/result、root/parent、pending/consumed IDs、audit IDs"],
        ["RevisionExecutionAuthorization", "actor/source/reason、最大Provider/Retrieval call、PENDING/CONSUMED、reservation"],
    ], [2300, 7060], font_size=8.7)
    add_heading(doc, "Revision状態遷移", 2)
    add_table(doc, ["順序", "State", "遷移条件"], [
        ["1", "IDLE", "通常実行または旧checkpoint読込"],
        ["2", "PLANNED / REQUESTED", "findingとbase Artifactを固定"],
        ["3A", "AUTHORIZATION_REQUIRED", "Safe Modeまたは有料内部Revision"],
        ["3B", "WAITING_UPSTREAM_RESULT", "隣接LayerへOutbox requestを保存"],
        ["4", "CONSUMING_REQUEST", "上流Layerがrequest相関とstaleを検証"],
        ["5", "EXECUTING", "budgetを消費し、対象と依存checkpointだけ実行"],
        ["6", "RESULT_READY", "finding dispositionと結果Artifactを保存"],
        ["7", "COMPLETED", "requesterが相関Resultを消費し再検証"],
        ["X", "BLOCKED", "budget枯渇、stale、duplicate conflict、Human Gate、未回答Provider"],
    ], [900, 2640, 5820], font_size=8.8)

    add_chapter(doc, 19, "Layer別Revision")
    add_table(doc, ["Layer", "Internal", "Upstream / Consumer"], [
        ["Producer", "Quality findingの対象Agent以降。General OpinionはRetrieval identityも分離", "Upstreamなし。Researcher Plan defectを受けResearch Plannerだけを修正"],
        ["Researcher", "Human REVISE planを別承認で実行。integrity repairは別", "ProducerへPlan修正。DeliberationのEvidence要求を処理してReportを返す"],
        ["Deliberation", "Primary/Counterargument/Manager/Reviewerの最小依存閉包", "ResearcherへEvidence要求。Conclusion requestを処理してDeliberation Resultを返す"],
        ["Conclusion", "Position/Evaluation/Integration/Reviewerの最小依存閉包", "Deliberationへ分析要求。Playwright requestはselection impactを判定して返す"],
        ["Playwright", "Narrative/Script/Citation/Visualの最小依存閉包", "Conclusionへ修正要求。Resultと新Handoffを検証して制作再開"],
    ], [1350, 3600, 4410], font_size=8.5)
    add_note(doc, "Routing rule", "Producer以外は直前Layerだけへ戻れます。新Evidenceが必要ならConclusion/Playwrightで捏造せず、隣接routeを連鎖させます。")

    add_chapter(doc, 20, "Recovery / Retry / Repair")
    add_table(doc, ["操作", "使う場面", "再実行規則"], [
        ["Resume", "正常なUpstream Revision Resultを受領", "相関Resultを消費し、stale依存checkpointだけ再計算"],
        ["Recover", "process/state/checkpoint障害", "保存Resultと予約を照合し、最後の未完了地点から再開"],
        ["Provider Retry", "課金済み可能性のある曖昧通信障害", "元reservationを保持し、別task identityで明示的一回"],
        ["Contract/Capability Repair", "同一model retryで直らないProvider契約", "異なる明示model、一回限り、互換bindingを監査保存"],
        ["Deterministic Repair", "一意に直せるmetadata/schema relation", "Provider/Retrieval 0、専用budget、前後hash、allowlist"],
    ], [1800, 3600, 3960], font_size=8.7)
    add_bullets(doc, [
        "未回答Provider requestが残る場合は実行有無を推測せず停止する。",
        "保存済みResult PMPが有効ならstate欠落を復元し、Providerを再呼出ししない。",
        "完了済みstageはlogical task IDとArtifact hashで再利用する。",
        "revision、retry、repair、retrieval reconstructionのbudget/authorizationを混同しない。",
    ])

    add_chapter(doc, 21, "Safe Mode")
    add_para(doc, "Demo Safe Modeは既定ONです。自動retryや追加有料Revisionを止めますが、request、plan、state、Outbox、auditの0-call保存は継続します。")
    add_table(doc, ["状況", "Safe Mode ON", "Safe Mode OFF"], [
        ["初回forward call", "明示Startなら実行", "実行"],
        ["Internal Revision", "authorization_requiredで停止し、reviseコマンド一回分だけ実行", "設定budget内で自動進行可能"],
        ["Upstream Request writer", "0-callなので保存", "保存"],
        ["Upstream consumerの有料処理", "明示revise/resume境界で実行", "設定budget内で実行"],
        ["Ambiguous provider retry", "自動再送しない。専用operator commandのみ", "同じく曖昧再送は自動化しない"],
    ], [2250, 3555, 3555], font_size=8.8)

    add_chapter(doc, 22, "CLI")
    add_para(doc, "`main.py`が唯一のCLI entry pointです。まず`--status`で停止層と推奨操作を確認し、Start/Resume/Recover/Retry/Reviseを使い分けます。")
    add_table(doc, ["目的", "コマンド"], [
        ["診断・状態", "--doctor、--status WORKFLOW_ID、--help、--version、--json、--verbose"],
        ["Mock検証", "--demo、--demo-full、--demo-e2e --provider mock"],
        ["Producer", "--producer-recover、--producer-revise、--producer-provider-retry、--producer-output-repair"],
        ["Researcher", "--researcher、--researcher-resume、--researcher-recover、--researcher-evidence、--researcher-accept、--researcher-accept-limitations、--researcher-revise、--researcher-revision-execute"],
        ["Deliberation", "--deliberation、--deliberation-resume、--deliberation-recover、--deliberation-revise、--deliberation-provider-retry"],
        ["Conclusion", "--conclusion、--conclusion-resume、--conclusion-recover、--conclusion-revise、--conclusion-provider-retry、--conclusion-select、--conclusion-integrate"],
        ["Playwright", "--playwright、--playwright-resume、--playwright-recover、--playwright-revise、--playwright-provider-retry"],
    ], [1900, 7460], font_size=8.4)
    add_code(doc, "py main.py --status <WORKFLOW_ID>\npy main.py --<layer>-revise <WORKFLOW_ID> --reason \"operator rationale\" --safe-mode\npy main.py --<layer>-recover <WORKFLOW_ID>")

    add_chapter(doc, 23, "Discord Control Plane")
    add_para(doc, "Discordは35コマンドを提供するControl Planeです。PMP transportではなく、Managerの開始、状態表示、Human Decision、Human Selection、Revision承認を呼び出します。Provider/Retrieval/Safe Modeは起動時設定を使い、コマンド単位overrideはありません。")
    add_table(doc, ["Layer", "代表コマンド"], [
        ["Producer", "!producer, !producer_topic, !producer_status, !producer_revise"],
        ["Researcher", "!researcher, !researcher_evidence, !researcher_accept, !researcher_accept_limitations, !researcher_revise, !researcher_revision_execute, !researcher_recover"],
        ["Deliberation", "!deliberation, !deliberation_status, !deliberation_result, !deliberation_resume, !deliberation_revise"],
        ["Conclusion", "!conclusion, !conclusion_options, !conclusion_select, !conclusion_integrate, !conclusion_resume, !conclusion_revise"],
        ["Playwright", "!playwright, !playwright_script, !playwright_citations, !playwright_visuals, !playwright_result, !playwright_resume, !playwright_revise"],
        ["Runtime", "!runtime_models [layer]"],
    ], [1800, 7560], font_size=8.5)
    add_note(doc, "Error handling", "完全TracebackはApplication Logへ保存し、Discordには700文字以内のoperational summaryを送ります。Layer実行失敗後は表示statusをERRORへ閉じ、RUNNINGを残しません。")

    add_chapter(doc, 24, "Delivery")
    add_para(doc, "Playwright Final Gate通過後、`storage/data/deliveries/<workflow_id>/`へ6ファイルを生成します。既存Delivery messageと成果物を検査し、restart/recoverで二重生成しません。")
    add_table(doc, ["ファイル", "内容"], [
        ["final_script_package.json", "全制作成果物とtraceabilityの機械可読package"],
        ["script.md", "最終台本"],
        ["citation_manifest.json", "Paragraph→Claim→Evidence→Source対応"],
        ["source_list.md", "使用Source一覧"],
        ["visual_plan.md", "Visual cueとasset/source指示"],
        ["production_notes.md", "制作条件、limitations、未解決事項"],
    ], [3000, 6360])

    add_chapter(doc, 25, "Logging / Audit / Traceability")
    add_bullets(doc, [
        "Message history: 全PMP envelopeと親子相関をLayer stateとJSONLへ保存。",
        "RD trace: Agent response metadataとManager stateにRD ID/version/hashを保存。",
        "Revision audit: request written/consumed、authorization、budget、reservation、result、blockedをevent化。",
        "Provider error: secretやraw全文を残さず、error class、hash、長さ、path、task/modelを保存。",
        "Provenance: evidence/source/analysis/claim/counterargument/change等を型別namespaceで検証。",
        "Human audit: Evidence DecisionとConclusion Selectionをcreate-once Artifact/PMPとして保存。",
    ])
    add_code(doc, "storage/data/logs/rd_access.jsonl\nstorage/data/artifacts/revision_audit/<workflow>/<request>/\nstorage/data/artifacts/human_evidence_decisions/<workflow>/\nstorage/data/artifacts/*_deterministic_repairs/<workflow>/")

    add_chapter(doc, 26, "Testing")
    add_para(doc, "検証は低コストから高コストへ進めます。実API固有でない問題をReal callで確認しません。")
    add_steps(doc, [
        "Static compileとdiff checkを行う。",
        "Unit、targeted regression、Schema、PMP、old checkpoint loadを実行する。",
        "Integration、Revision roundtrip、Fault Injection、Mock Recoveryを実行する。",
        "31 RD STRICT、22 Structured Output Schema、Doctorを実行する。",
        "一時storageで5 Layer Mock E2EとDelivery 6ファイルを確認する。",
        "保存済みworkflowを一時領域へcloneし、元hash不変・API 0で互換性を確認する。",
    ])
    add_code(doc, "py scripts/verify.py\npy -m unittest tests.unit.test_revision_contract -v\npy scripts/verify_revision_architecture_production_clone.py --source-project C:\\Projects\\PRDCP_v2 --workflow-id <WORKFLOW_ID>")

    add_chapter(doc, 27, "新Agent追加方法")
    add_steps(doc, [
        "Agent ID、責務、入力/出力、禁止事項、timeout、modelを設計する。",
        "Pydantic input/output modelを追加し、Structured Outputなら22-root監査相当へ登録する。",
        "`role_definitions/<layer>/`へRDを追加しregistry/model設定へ登録する。",
        "共通Agent execution pipelineを使うRegistry実装へ追加する。",
        "Managerの依存順、checkpoint、task identity、PMP sender/receiverを更新する。",
        "Mock fixture、schema test、RD STRICT、provider reservation、recovery/idempotency testを追加する。",
        "README/Developer Guide/Doctorの件数と説明を更新する。",
    ])
    add_note(doc, "禁止", "Layer固有の独自retry、自由形式Structured Output dict、非決定的task ID、未登録PMP typeを追加しないでください。")

    add_chapter(doc, 28, "新Layer / Revision route追加時の注意")
    add_steps(doc, [
        "canonical ownerを決め、producer→persistence→direct consumer→transitive downstream→Deliveryを変更前後で監査する。",
        "LayerIdと隣接`UPSTREAM_LAYER`を更新し、Layer skipping禁止を維持する。",
        "Revision Request/ResultのArtifact type、finding disposition、Human impactを定義する。",
        "request単位Outbox、consumer、result writer、downstream resumeを完全往復で実装する。",
        "internal/upstream budget、authorization、Provider/Retrieval reservationを分離する。",
        "stale/duplicate/restart/fault/old checkpoint/Delivery exactly-onceテストを追加する。",
        "CLI/Discord/README/Doctor/specification registryを同じcommitで同期する。",
    ])
    add_note(doc, "Contract drift", "共有Schema変更はdirect consumerだけでなく、persist→reload→全downstream validationまで試験してください。writer-only routeを完成扱いしないでください。", color="FFF4CE")

    add_chapter(doc, 29, "Known limitations")
    add_bullets(doc, [
        "Mock E2Eは制御系・Schema・保存・接続を検証しますが、実OpenRouterの応答品質や外部可用性は保証しません。",
        "OpenRouter model/endpoint metadataは変化するため、実運用直前に`--doctor`で再確認が必要です。",
        "Discordは技術的なprovider retryやcontract/capability repairの全操作を公開せず、一部はCLI専用です。",
        "古いcheckpointはread adapterで互換化しますが、未知・曖昧な破損を推測修復しません。",
        "Human Evidenceの受容は不足情報をEvidenceへ変換せず、limitationsとして残します。",
        "Conclusionの意味変更にはHuman reselectionが必要で、無人E2Eでは自動迂回できません。",
        "file Outboxは単一ホストのcanonical storage設計であり、分散message brokerの配送保証を提供しません。",
    ])

    add_chapter(doc, 30, "Troubleshooting")
    add_table(doc, ["症状", "確認", "正しい操作"], [
        ["Schema 400", "22 root strict audit、dynamic specialization、free-form dict", "Schemaを明示model化。strictを弱めない"],
        ["Provider応答不明", "request/error PMP、reservation、authorization", "自動再送せずprovider-retryを一回だけ"],
        ["WAITING_UPSTREAM_REVISION", "Outbox requestと期待Result ID/hash", "上流をrevise/resume後、下流resume"],
        ["AUTHORIZATION_REQUIRED", "Safe Mode、request、budget、最大call数", "対象Layerのrevise/revision-executeを明示"],
        ["Stale result", "base/result Artifact hash、workflow/request/message", "新しいrequestを作る。古いResultを適用しない"],
        ["Human Gate停止", "Evidence gap/hard finding、候補、limitations", "人間がDecision/Selectionを記録"],
        ["Deliveryなし", "Final Gate finding、Manifest、selection、既存Delivery", "deterministic repair→internal revision→upstreamの順"],
        ["Discord RUNNING残り", "Application Logの完全Traceback", "現行error handlerがERRORへ閉じることを確認"],
    ], [1800, 3360, 4200], font_size=8.2)
    add_heading(doc, "初動チェック", 2)
    add_code(doc, "py main.py --doctor\npy main.py --status <WORKFLOW_ID>\npy main.py --status <WORKFLOW_ID> --json\npy scripts/verify.py")
    add_note(doc, "Escalation", "認証、残高、外部Provider障害、大量Real API、破壊的migration、既存workflow削除、Human Decisionが必要な場合は停止し、operatorへ状況と選択肢を報告します。")

    add_heading(doc, "Source of Truth", 1)
    add_table(doc, ["対象", "正本パス"], [
        ["Runtime", "main.py, runtime.py, cli_app/, discord_app/"],
        ["Common contracts", "common/models/, common/validation/, specifications/common/"],
        ["Layer workflows", "producer/, researcher/, deliberation/, conclusion/, playwright/"],
        ["RD / Prompt", "role_definitions/, common/role_definitions/, common/prompting/"],
        ["Provider / Retrieval", "providers/, retrieval/, common/structured_outputs.py"],
        ["Persistence", "storage/, storage/data/ (runtime data; source control対象外を原則とする)"],
        ["Tests", "tests/unit/, tests/integration/, scripts/verify.py"],
    ], [2500, 6860])
    add_para(doc, "本書とREADMEに差異がある場合は、active code、Pydantic model、specification registry、自動テストの順に確認し、文書を同じ変更で更新してください。")

    for section in doc.sections:
        configure_section(section)
        add_header_footer(section, first=section is doc.sections[0])

    core = doc.core_properties
    core.title = "PRDCP v2 Developer Guide"
    core.subject = "Five-layer architecture, revision, recovery, and operations"
    core.author = "PRDCP Project"
    core.keywords = "PRDCP, PMP, Revision, Structured Output, Recovery"
    core.comments = "Generated from the active PRDCP v2 implementation contract."
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
